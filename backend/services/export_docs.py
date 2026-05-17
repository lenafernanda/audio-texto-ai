"""Geração de PDF e DOCX para exportação."""

from __future__ import annotations

import io
from typing import Any


def build_txt(texto: str, resumo: str, estudo: str) -> bytes:
    body = [
        "TRANSCRIÇÃO",
        "=" * 40,
        texto,
        "",
        "RESUMO",
        "=" * 40,
        resumo,
        "",
        "VERSÃO PARA ESTUDO",
        "=" * 40,
        estudo,
    ]
    return "\n".join(body).encode("utf-8")


def build_pdf(texto: str, resumo: str, estudo: str, titulo: str = "Transcrição") -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, titulo)
    pdf.ln(4)

    sections = [
        ("Transcrição", texto),
        ("Resumo", resumo),
        ("Versão para estudo", estudo),
    ]
    for heading, content in sections:
        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(0, 7, heading)
        pdf.set_font("Helvetica", size=10)
        safe = _latin1_safe(content or "(vazio)")
        pdf.multi_cell(0, 5, safe)
        pdf.ln(3)

    out = pdf.output()
    return out if isinstance(out, bytes) else out.encode("latin-1", errors="replace")


def build_docx(texto: str, resumo: str, estudo: str, titulo: str = "Transcrição") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(titulo, level=0)
    doc.add_heading("Transcrição", level=1)
    for para in (texto or "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_heading("Resumo", level=1)
    doc.add_paragraph(resumo or "")
    doc.add_heading("Versão para estudo", level=1)
    for line in (estudo or "").splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _latin1_safe(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")
